import io
import re
import csv
import logging
from datetime import datetime, timezone
from typing import List, Optional

from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, BackgroundTasks
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, update, delete

from database import get_db, AsyncSessionLocal
from auth import get_current_user
from orm_models import Document, Questionnaire
from ai_service import answer_questions_with_docs

logger = logging.getLogger(__name__)
router = APIRouter()

MAX_QUESTIONS = 100
MAX_FILE_SIZE = 5 * 1024 * 1024


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def extract_text(file_bytes: bytes, ext: str) -> str:
    if ext == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in reader.pages:
            text += (page.extract_text() or "") + "\n"
        return text.strip()
    return file_bytes.decode("utf-8", errors="ignore").strip()


def _is_title_or_heading(text: str) -> bool:
    t = text.strip()
    if not t.endswith("?") and (t.isupper() or (len(t.split()) <= 6 and t.istitle())):
        return True
    heading_words = {"questionnaire", "form", "assessment", "survey", "checklist", "template", "document", "policy"}
    lower_words = set(t.lower().split())
    if not t.endswith("?") and len(t.split()) <= 8 and heading_words & lower_words:
        return True
    return False


def parse_questions(text: str) -> List[str]:
    lines = text.split("\n")

    q_num_pattern = re.compile(r"^\s*Q(\d+)[\.\):\s]\s*(.+)", re.IGNORECASE)
    q_style, current = [], ""
    for line in lines:
        stripped = line.strip()
        m = q_num_pattern.match(stripped)
        if m:
            if current:
                q_style.append(current.strip())
            current = m.group(2).strip()
        elif current and stripped:
            current += " " + stripped
        elif not stripped and current:
            q_style.append(current.strip())
            current = ""
    if current:
        q_style.append(current.strip())
    if q_style:
        return [q for q in q_style if len(q) > 5][:MAX_QUESTIONS]

    num_pattern = re.compile(r"^\s*\d+[\.\):\s]\s*(.+)", re.IGNORECASE)
    num_style, current = [], ""
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current:
                full = current.strip()
                if full.endswith("?") and not _is_title_or_heading(full):
                    num_style.append(full)
                current = ""
            continue
        m = num_pattern.match(stripped)
        if m:
            if current:
                full = current.strip()
                if full.endswith("?") and not _is_title_or_heading(full):
                    num_style.append(full)
            current = m.group(1).strip()
        elif current:
            current += " " + stripped
    if current:
        full = current.strip()
        if full.endswith("?") and not _is_title_or_heading(full):
            num_style.append(full)
    if num_style:
        return num_style[:MAX_QUESTIONS]

    return [l.strip() for l in lines if l.strip().endswith("?") and len(l.strip()) > 10 and not _is_title_or_heading(l.strip())][:MAX_QUESTIONS]


def _serialize_for_json(obj):
    """Recursively convert datetime objects to ISO strings for JSONB storage."""
    if isinstance(obj, datetime):
        return obj.isoformat()
    if isinstance(obj, dict):
        return {k: _serialize_for_json(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_serialize_for_json(item) for item in obj]
    return obj


def _version_summary(v: dict) -> dict:
    ca = v.get("created_at")
    return {
        "version_number": v.get("version_number"),
        "created_at": ca.isoformat() if hasattr(ca, "isoformat") else ca,
        "answers_found_count": v.get("answers_found_count", 0),
        "confidence_counts": v.get("confidence_counts", {}),
        "answers": v.get("answers", []),
    }


def _conf_counts(answers: list) -> dict:
    cc = {"HIGH": 0, "MEDIUM": 0, "LOW": 0}
    for a in answers:
        c = (a.get("confidence") or "LOW").upper()
        if c in cc:
            cc[c] += 1
    return cc


def _q_to_dict(q: Questionnaire) -> dict:
    """Convert ORM Questionnaire to plain dict for use in build_docx and responses."""
    return {
        "id": q.id,
        "name": q.name,
        "status": q.status,
        "questions": q.questions or [],
        "answers": q.answers or [],
        "versions": q.versions or [],
        "document_ids": q.document_ids or [],
        "question_count": q.question_count,
        "created_at": q.created_at.isoformat() if q.created_at else None,
        "completed_at": q.completed_at.isoformat() if q.completed_at else None,
        "error_message": q.error_message,
    }


def build_docx(q: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    answers = q.get("answers", [])
    found_count = sum(1 for a in answers if a.get("found"))

    doc.add_heading(q.get("name", "Questionnaire Report"), 0)
    subtitle = doc.add_paragraph("Generated by TrustBase AI")
    subtitle.runs[0].italic = True
    subtitle.runs[0].font.color.rgb = RGBColor(0x38, 0xBD, 0xF8)

    ca = q.get("completed_at")
    if ca:
        try:
            date_str = datetime.fromisoformat(str(ca)).strftime("%B %d, %Y")
        except Exception:
            date_str = str(ca)[:10]
    else:
        date_str = datetime.now().strftime("%B %d, %Y")

    doc.add_paragraph(f"Date Generated: {date_str}")
    doc.add_paragraph(f"Total Questions: {len(answers)}")
    doc.add_paragraph(f"Answered with Citations: {found_count}")
    doc.add_paragraph(f"Not Found in References: {len(answers) - found_count}")
    cc = _conf_counts(answers)
    doc.add_paragraph(f"Confidence — High: {cc['HIGH']}  Medium: {cc['MEDIUM']}  Low: {cc['LOW']}")
    doc.add_paragraph("")

    CONF_COLORS = {
        "HIGH": RGBColor(0x16, 0xA3, 0x4A),
        "MEDIUM": RGBColor(0xD9, 0x77, 0x06),
        "LOW": RGBColor(0xDC, 0x26, 0x26),
    }

    for i, ans in enumerate(answers, 1):
        doc.add_heading(f"Q{i}. {ans.get('question', '')}", level=2)
        if ans.get("found"):
            doc.add_paragraph(ans.get("answer", ""))
            conf = (ans.get("confidence") or "LOW").upper()
            cp = doc.add_paragraph()
            run = cp.add_run(f"Confidence: {conf}")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = CONF_COLORS.get(conf, CONF_COLORS["LOW"])
            src = ans.get("source_document", "")
            cit = ans.get("citation", "")
            ev = ans.get("evidence_text", "")
            if src:
                sp = doc.add_paragraph()
                sp.add_run(f"Source: {src}").italic = True
                if cit:
                    sp.add_run(f' — "{cit}"').italic = True
            if ev and ev != cit:
                ep = doc.add_paragraph(f'Evidence: "{ev}"')
                ep.runs[0].font.size = Pt(10)
                ep.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        else:
            p = doc.add_paragraph("Not found in references.")
            p.runs[0].italic = True
            p.runs[0].font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
        if ans.get("is_edited"):
            ep = doc.add_paragraph("[Manually edited]")
            ep.runs[0].font.size = Pt(9)
            ep.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Request bodies
# ---------------------------------------------------------------------------

class ProcessRequest(BaseModel):
    document_ids: Optional[List[str]] = None


class EditAnswerRequest(BaseModel):
    answer: str


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@router.get("")
async def list_questionnaires(
    user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Questionnaire)
        .where(Questionnaire.user_id == user_id)
        .order_by(Questionnaire.created_at.desc())
    )
    qs = result.scalars().all()
    return [
        {
            "id": q.id,
            "name": q.name,
            "status": q.status,
            "question_count": q.question_count,
            "created_at": q.created_at.isoformat() if q.created_at else None,
            "completed_at": q.completed_at.isoformat() if q.completed_at else None,
            "error_message": q.error_message,
        }
        for q in qs
    ]


@router.post("")
async def upload_questionnaire(
    file: UploadFile = File(...),
    user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    file_bytes = await file.read()
    if len(file_bytes) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File too large (max 5 MB)")

    fname = file.filename or "questionnaire"
    ext = fname.rsplit(".", 1)[-1].lower() if "." in fname else "txt"
    if ext not in ("pdf", "txt", "text", "md"):
        raise HTTPException(status_code=400, detail="Only PDF and TXT files are supported")

    try:
        text = extract_text(file_bytes, ext)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to extract text: {e}")

    if not text:
        raise HTTPException(status_code=400, detail="Could not extract text from file")

    questions = parse_questions(text)
    if not questions:
        raise HTTPException(status_code=400, detail="No questions found in file")

    q_name = fname.rsplit(".", 1)[0] if "." in fname else fname
    q = Questionnaire(
        user_id=user_id,
        name=q_name,
        status="pending",
        questions=questions,
        question_count=len(questions),
        answers=[],
        versions=[],
        document_ids=[],
    )
    db.add(q)
    await db.commit()
    await db.refresh(q)

    return {
        "id": q.id,
        "name": q.name,
        "status": q.status,
        "questions": q.questions,
        "question_count": q.question_count,
    }


@router.get("/{q_id}")
async def get_questionnaire(
    q_id: str,
    user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Questionnaire).where(Questionnaire.id == q_id, Questionnaire.user_id == user_id)
    )
    q = result.scalar_one_or_none()
    if not q:
        raise HTTPException(status_code=404, detail="Not found")
    d = _q_to_dict(q)
    d["versions"] = [_version_summary(v) for v in (q.versions or [])]
    return d


@router.delete("/{q_id}")
async def delete_questionnaire(
    q_id: str,
    user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Questionnaire).where(Questionnaire.id == q_id, Questionnaire.user_id == user_id)
    )
    q = result.scalar_one_or_none()
    if not q:
        raise HTTPException(status_code=404, detail="Not found")
    await db.delete(q)
    await db.commit()
    return {"success": True}


@router.post("/{q_id}/process")
async def process_questionnaire(
    q_id: str,
    body: ProcessRequest,
    background_tasks: BackgroundTasks,
    user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Questionnaire).where(Questionnaire.id == q_id, Questionnaire.user_id == user_id)
    )
    q = result.scalar_one_or_none()
    if not q:
        raise HTTPException(status_code=404, detail="Not found")
    if q.status == "processing":
        raise HTTPException(status_code=400, detail="Already processing")

    doc_ids = body.document_ids or []
    if not doc_ids:
        docs_result = await db.execute(
            select(Document.id).where(Document.user_id == user_id)
        )
        doc_ids = [str(row[0]) for row in docs_result.all()]
    if not doc_ids:
        raise HTTPException(status_code=400, detail="No reference documents available. Please upload documents first.")

    await db.execute(
        update(Questionnaire)
        .where(Questionnaire.id == q_id)
        .values(status="processing", document_ids=doc_ids)
    )
    await db.commit()

    background_tasks.add_task(_bg_process, q_id, doc_ids)
    return {"status": "processing"}


@router.post("/{q_id}/regenerate")
async def regenerate_questionnaire(
    q_id: str,
    background_tasks: BackgroundTasks,
    user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Questionnaire).where(Questionnaire.id == q_id, Questionnaire.user_id == user_id)
    )
    q = result.scalar_one_or_none()
    if not q:
        raise HTTPException(status_code=404, detail="Not found")
    if q.status == "processing":
        raise HTTPException(status_code=400, detail="Already processing")

    docs_result = await db.execute(
        select(Document.id).where(Document.user_id == user_id)
    )
    doc_ids = [str(row[0]) for row in docs_result.all()]
    if not doc_ids:
        raise HTTPException(status_code=400, detail="No reference documents available. Please upload documents first.")

    # Save current answers as a version snapshot
    current_answers = q.answers or []
    if current_answers:
        versions = list(q.versions or [])
        version_num = len(versions) + 1
        found_ct = sum(1 for a in current_answers if a.get("found"))
        new_version = _serialize_for_json({
            "version_number": version_num,
            "created_at": datetime.now(timezone.utc),
            "answers": current_answers,
            "answers_found_count": found_ct,
            "confidence_counts": _conf_counts(current_answers),
        })
        versions.append(new_version)
        await db.execute(
            update(Questionnaire)
            .where(Questionnaire.id == q_id)
            .values(versions=versions)
        )

    await db.execute(
        update(Questionnaire)
        .where(Questionnaire.id == q_id)
        .values(status="processing", document_ids=doc_ids, answers=[], error_message=None)
    )
    await db.commit()

    background_tasks.add_task(_bg_process, q_id, doc_ids)
    return {"status": "processing", "message": "Regeneration started"}


@router.patch("/{q_id}/answers/{idx}")
async def edit_answer(
    q_id: str,
    idx: int,
    body: EditAnswerRequest,
    user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Questionnaire).where(Questionnaire.id == q_id, Questionnaire.user_id == user_id)
    )
    q = result.scalar_one_or_none()
    if not q:
        raise HTTPException(status_code=404, detail="Not found")
    answers = list(q.answers or [])
    if idx < 0 or idx >= len(answers):
        raise HTTPException(status_code=400, detail="Invalid answer index")

    updated_answer = dict(answers[idx])
    updated_answer["answer"] = body.answer
    updated_answer["is_edited"] = True
    answers[idx] = updated_answer

    await db.execute(
        update(Questionnaire).where(Questionnaire.id == q_id).values(answers=answers)
    )
    await db.commit()
    return updated_answer


@router.post("/{q_id}/answers/{idx}/regenerate")
async def regenerate_single_answer(
    q_id: str,
    idx: int,
    user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Questionnaire).where(Questionnaire.id == q_id, Questionnaire.user_id == user_id)
    )
    q = result.scalar_one_or_none()
    if not q:
        raise HTTPException(status_code=404, detail="Not found")
    questions = q.questions or []
    if idx < 0 or idx >= len(questions):
        raise HTTPException(status_code=400, detail="Invalid question index")

    docs_result = await db.execute(
        select(Document).where(Document.user_id == user_id)
    )
    all_docs = docs_result.scalars().all()
    if not all_docs:
        raise HTTPException(status_code=400, detail="No reference documents available")

    docs = [{"name": d.name, "content": d.content} for d in all_docs]
    ai_answers = await answer_questions_with_docs([questions[idx]], docs)
    new_answer = ai_answers[0] if ai_answers else {
        "question": questions[idx], "answer": "Processing error", "found": False,
        "source_document": None, "citation": None, "confidence": "LOW",
        "evidence_text": None, "is_edited": False,
    }

    answers = list(q.answers or [])
    if idx < len(answers):
        answers[idx] = new_answer
    else:
        while len(answers) < idx:
            answers.append({})
        answers.append(new_answer)

    await db.execute(
        update(Questionnaire).where(Questionnaire.id == q_id).values(answers=answers)
    )
    await db.commit()
    return new_answer


@router.get("/{q_id}/download-docx")
async def download_docx(
    q_id: str,
    user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Questionnaire).where(Questionnaire.id == q_id, Questionnaire.user_id == user_id)
    )
    q = result.scalar_one_or_none()
    if not q:
        raise HTTPException(status_code=404, detail="Not found")
    if q.status != "completed":
        raise HTTPException(status_code=400, detail="Questionnaire not yet completed")

    docx_bytes = build_docx(_q_to_dict(q))
    safe_name = re.sub(r"[^\w\-]", "_", q.name)
    return StreamingResponse(
        iter([docx_bytes]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}_report.docx"'},
    )


@router.get("/{q_id}/download")
async def download_csv(
    q_id: str,
    user_id: str = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Questionnaire).where(Questionnaire.id == q_id, Questionnaire.user_id == user_id)
    )
    q = result.scalar_one_or_none()
    if not q:
        raise HTTPException(status_code=404, detail="Not found")
    if q.status != "completed":
        raise HTTPException(status_code=400, detail="Questionnaire not yet completed")

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["#", "Question", "Answer", "Found", "Confidence", "Source Document", "Citation"])
    for i, ans in enumerate(q.answers or [], 1):
        writer.writerow([
            i, ans.get("question", ""), ans.get("answer", ""),
            "Yes" if ans.get("found") else "No",
            ans.get("confidence", ""),
            ans.get("source_document") or "",
            ans.get("citation") or "",
        ])
    safe_name = re.sub(r"[^\w\-]", "_", q.name)
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}_answers.csv"'},
    )


# ---------------------------------------------------------------------------
# Background task
# ---------------------------------------------------------------------------

async def _bg_process(q_id: str, doc_ids: List[str]):
    async with AsyncSessionLocal() as db:
        try:
            result = await db.execute(select(Questionnaire).where(Questionnaire.id == q_id))
            q = result.scalar_one_or_none()
            if not q:
                return

            questions = q.questions or []

            # Fetch documents
            docs_result = await db.execute(
                select(Document).where(Document.id.in_(doc_ids))
            )
            doc_rows = docs_result.scalars().all()
            docs = [{"name": d.name, "content": d.content} for d in doc_rows]

            if not docs:
                await db.execute(
                    update(Questionnaire)
                    .where(Questionnaire.id == q_id)
                    .values(status="failed", error_message="Reference documents not found")
                )
                await db.commit()
                return

            answers = await answer_questions_with_docs(questions, docs)

            await db.execute(
                update(Questionnaire)
                .where(Questionnaire.id == q_id)
                .values(
                    status="completed",
                    answers=answers,
                    completed_at=datetime.now(timezone.utc),
                    error_message=None,
                )
            )
            await db.commit()

        except Exception as e:
            logger.error(f"Background processing error for {q_id}: {e}")
            try:
                await db.execute(
                    update(Questionnaire)
                    .where(Questionnaire.id == q_id)
                    .values(status="failed", error_message=str(e))
                )
                await db.commit()
            except Exception:
                pass
